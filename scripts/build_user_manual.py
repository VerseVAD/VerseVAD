"""Build the maintainable VerseVAD Word user manual from its Markdown source."""

from __future__ import annotations

import re
import tomllib
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "VerseVAD_User_Manual_Source.md"
OUTPUT = ROOT / "docs" / "VerseVAD_User_Manual.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
TABLE_CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "193247"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TERRACOTTA = "B14E2C"
MUTED = "66727C"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D1D8"
WHITE = "FFFFFF"


def _rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float = 1.0,
) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.widow_control = True


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in TABLE_CELL_MARGINS.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _set_table_borders(table, color: str = BORDER, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _column_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    if columns == 1:
        return [CONTENT_WIDTH_DXA]
    if columns == 2:
        first_lengths = [len(row[0]) for row in rows]
        if max(first_lengths, default=0) <= 28:
            return [2700, CONTENT_WIDTH_DXA - 2700]
        return [3600, CONTENT_WIDTH_DXA - 3600]
    if columns == 3:
        return [2250, 3150, 3960]
    if columns == 4:
        return [2100, 3000, 1500, 2760]
    base = CONTENT_WIDTH_DXA // columns
    widths = [base] * columns
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    _set_table_geometry(table, _column_widths(rows))
    _set_table_borders(table)
    for row_index, values in enumerate(rows):
        _prevent_row_split(table.rows[row_index])
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _set_spacing(paragraph, after=0, line=1.05)
            if row_index == 0:
                paragraph.paragraph_format.keep_with_next = True
            _add_inline(paragraph, value, table_text=True)
            if row_index == 0:
                _shade_cell(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = _rgb(NAVY)
        if row_index == 0:
            _repeat_header(table.rows[0])
    spacer = document.add_paragraph()
    _set_spacing(spacer, after=3, line=1.0)


def _add_callout(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    _set_spacing(paragraph, before=4, after=6, line=1.15)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge_name in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "5")
        edge.set(qn("w:space"), "5")
        edge.set(qn("w:color"), "D7DEE5")
        borders.append(edge)
    p_pr.append(borders)
    _add_inline(paragraph, text)
    for run in paragraph.runs:
        run.font.color.rgb = _rgb(NAVY)


def _add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph(style="Formula")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        _set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def _add_inline(paragraph, text: str, *, table_text: bool = False) -> None:
    size = 9.3 if table_text else 11
    pieces = INLINE_PATTERN.split(text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            _set_run_font(run, size=size, bold=True)
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            _set_run_font(run, name="Consolas", size=8.8 if table_text else 9.5, color=DARK_BLUE)
        elif piece.startswith("*") and piece.endswith("*"):
            run = paragraph.add_run(piece[1:-1])
            _set_run_font(run, size=size, italic=True)
        else:
            segments = piece.split("\n")
            for index, segment in enumerate(segments):
                if index:
                    paragraph.add_run().add_break()
                if segment:
                    run = paragraph.add_run(segment)
                    _set_run_font(run, size=size)


def _add_numbering_definition(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    first_number = numbering.find(qn("w:num"))
    if first_number is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_number), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if not bullet:
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)


def _add_list_item(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    _apply_numbering(paragraph, num_id)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    _set_spacing(paragraph, after=4, line=1.25)
    _add_inline(paragraph, text)


def _add_field(paragraph, instruction: str, placeholder: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.append(begin)
    run.append(instr)
    run.append(separate)
    run.append(text)
    run.append(end)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = _rgb("202A33")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Manual Title" not in styles:
        title = styles.add_style("Manual Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title = styles["Manual Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = _rgb(NAVY)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(118)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True

    if "Manual Subtitle" not in styles:
        subtitle = styles.add_style("Manual Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        subtitle = styles["Manual Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = _rgb(TERRACOTTA)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(30)
    subtitle.paragraph_format.keep_with_next = True

    if "Formula" not in styles:
        formula = styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    else:
        formula = styles["Formula"]
    formula.font.name = "Consolas"
    formula._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    formula._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    formula.font.size = Pt(9.5)
    formula.font.color.rgb = _rgb(DARK_BLUE)
    formula.paragraph_format.left_indent = Inches(0.18)
    formula.paragraph_format.space_before = Pt(2)
    formula.paragraph_format.space_after = Pt(8)
    formula.paragraph_format.line_spacing = 1.1


def _configure_page(
    document: Document,
    version: str,
    *,
    header_title: str = "VerseVAD User Manual",
) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(paragraph, after=0, line=1.0)
    run = paragraph.add_run(header_title)
    _set_run_font(run, size=9, color=MUTED, bold=True)
    run = paragraph.add_run(f"  |  {version}")
    _set_run_font(run, size=9, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_spacing(paragraph, after=0, line=1.0)
    run = paragraph.add_run("Page ")
    _set_run_font(run, size=9, color=MUTED)
    _add_field(paragraph, "PAGE", "1")


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def build_document_from_source(
    *,
    source: Path,
    output: Path,
    title: str,
    subject: str,
    header_title: str,
    comments: str,
) -> Path:
    project_data = tomllib.loads((ROOT / "pyproject.toml").read_text(encoding="utf-8"))
    version = project_data["project"]["version"]
    try:
        updated = date.today().strftime("%B %-d, %Y")
    except ValueError:
        updated = date.today().strftime("%B %#d, %Y")
    markdown = source.read_text(encoding="utf-8")
    markdown = markdown.replace("{{VERSION}}", version).replace("{{DATE}}", updated)

    document = Document()
    _configure_styles(document)
    _configure_page(document, version, header_title=header_title)
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.author = "VerseVAD"
    document.core_properties.keywords = (
        "VerseVAD, valence, arousal, dominance, affective lexicon, "
        "concreteness, corpus"
    )
    document.core_properties.comments = comments

    bullet_num_id = _add_numbering_definition(document, bullet=True)
    numbered_num_id: int | None = None
    in_numbered_list = False
    cover_mode = True
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        combined = ""
        for line in paragraph_buffer:
            hard_break = line.endswith("  ")
            content = line.rstrip()
            if combined:
                combined += "\n" if previous_hard_break else " "
            combined += content
            previous_hard_break = hard_break
        paragraph_buffer = []
        paragraph = document.add_paragraph()
        if cover_mode:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_spacing(paragraph, after=4, line=1.1)
            _add_inline(paragraph, combined)
            for run in paragraph.runs:
                run.font.color.rgb = _rgb(MUTED)
                run.font.size = Pt(10.5)
        elif combined.startswith("`") and combined.endswith("`"):
            paragraph.style = "Formula"
            run = paragraph.add_run(combined[1:-1])
            _set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            _set_spacing(paragraph, after=6, line=1.25)
            _add_inline(paragraph, combined)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped == "[[PAGEBREAK]]":
            flush_paragraph()
            document.add_page_break()
            cover_mode = False
            in_numbered_list = False
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ValueError(f"Unclosed fenced code block in {source}")
            _add_code_block(document, code_lines)
            in_numbered_list = False
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = _parse_table(table_lines)
            if rows:
                _add_table(document, rows)
            in_numbered_list = False
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2)
            if cover_mode and level == 1:
                paragraph = document.add_paragraph(style="Manual Title")
                _add_inline(paragraph, text)
            elif cover_mode and level == 2:
                paragraph = document.add_paragraph(style="Manual Subtitle")
                _add_inline(paragraph, text)
            else:
                paragraph = document.add_paragraph(style=f"Heading {level}")
                _add_inline(paragraph, text)
            in_numbered_list = False
            index += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("> "):
                quote_lines.append(lines[index].strip()[2:])
                index += 1
            _add_callout(document, " ".join(quote_lines))
            in_numbered_list = False
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            flush_paragraph()
            if not in_numbered_list:
                numbered_num_id = _add_numbering_definition(document, bullet=False)
                in_numbered_list = True
            assert numbered_num_id is not None
            _add_list_item(document, numbered.group(1), numbered_num_id)
            index += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            _add_list_item(document, stripped[2:], bullet_num_id)
            in_numbered_list = False
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            in_numbered_list = False
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def build_manual() -> Path:
    return build_document_from_source(
        source=SOURCE,
        output=OUTPUT,
        title="VerseVAD User Manual",
        subject="Local lexical-evidence analysis user manual",
        header_title="VerseVAD User Manual",
        comments="Generated from docs/VerseVAD_User_Manual_Source.md",
    )


if __name__ == "__main__":
    print(build_manual())
